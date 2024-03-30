
import contextlib
import tempfile
import shutil
import os

# Define the context manager in the main script
@contextlib.contextmanager
def temp_document_context(temp_file_paths):
    try:
        # Yield the list of temporary file paths back to the caller to use
        yield temp_file_paths
    finally:
        # Cleanup: Delete temporary files
        for path in temp_file_paths:
            os.unlink(path)


import io
import re
import gc
import json
import streamlit as st
import subprocess
from document_processing import process_documents, build_inverted_index
from search import search
from llmrouter import LLMRouter
from docx import Document
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.pagesizes import A4
from reportlab.platypus import Paragraph, Spacer, SimpleDocTemplate
from reportlab.lib.styles import getSampleStyleSheet
import markdown
from markdown.extensions import Extension
from markdown.treeprocessors import Treeprocessor
from dotenv import load_dotenv
from bs4 import BeautifulSoup

# Load environment variables from .env file
load_dotenv()

# Setup LLMRouter
llm_router = LLMRouter(
   anthropic_api_key=os.environ.get("ANTHROPIC_API_KEY"),
   openai_api_key=os.environ.get("OPENAI_API_KEY"),
   together_api_key=os.environ.get("TOGETHER_API_KEY")
)

def generate_index_and_abstract(instruction, length, language):
    messages = [
        {"role": "user", "content": f"Generate an index with points and subpoints, as well as an abstract for an essay based on the following instruction: {instruction}. The length should be {length}."}
    ]
    
    index_and_abstract = llm_router.generate("claude-3-opus-20240229", messages, max_tokens=2000, temperature=0.95, top_p=1.0, system=f"You are an expert professor with a formed mind and opinions, able to understand and write about complex topics in an academic manner with technical language in perfect {language}, as would be seen from a doctorate. The main sections should be numbered with Roman numerals, the subsections with letters and the subsubsections with Arabic numerals. When introducing the index of the abstract it must be put like 'Index:' / 'Índice:', 'Resumen:' / 'Abstract:' and etc. You must not mention anything outside of the paper/essay itself nor any metarreferencing.")
    
    messages.append({"role": "assistant", "content": index_and_abstract})
    messages.append({"role": "user", "content": "Improve the index and abstract, making sure it is innovative, complete, and well-structured."})
    
    improved_index_and_abstract = llm_router.generate("claude-3-opus-20240229", messages, max_tokens=2000, temperature=0.95, top_p=1.0, system=f"You are an expert professor with a formed mind and opinions, able to understand and write about complex topics in an academic manner with technical language in perfect {language}, as would be seen from a doctorate. The main sections should be numbered with Roman numerals, the subsections with letters and the subsubsections with Arabic numerals. You must not mention anything outside of the paper/essay itself nor any metarreferencing.")
    
    return improved_index_and_abstract

def download_papers(queries):
    downloaded_queries = {}
    if os.path.exists("downloaded_queries.json"):
        with open("downloaded_queries.json", "r") as file:
            downloaded_queries = json.load(file)
    
    for query in queries.split(","):
        if not query.strip():
            continue
        if query in downloaded_queries:
            print(f"Skipping query '{query}' as it has already been downloaded.")
        else:
            subprocess.run(["python", "-m", "PyPaperBot", f"--query={query}", "--scholar-pages=1", "--dwn-dir=data", "--max-dwn-cites=1"])
            downloaded_queries[query] = "data"
    
    with open("downloaded_queries.json", "w") as file:
        json.dump(downloaded_queries, file)

def process_downloaded_papers():
    folder_path = "data"
    documents = process_documents(folder_path)
    print("Documents processed")
    gc.collect()
    inverted_index = build_inverted_index(documents)
    gc.collect()
    print("Index builded")
    return documents, inverted_index

def extract_sections(text):
    sections = {}
    # Adjusted to include Spanish and French translations for "Title", "Abstract", and "Index"
    section_pattern = re.compile(
        r"(Title|Título|Titre|Abstract|Resumen|Résumé|Index|Índice|Indice):\s*(.*?)(?=\n\n[A-Z]|$)", 
        re.DOTALL
    )
    for match in section_pattern.finditer(text):
        section_name = match.group(1).lower()
        # Standardize the keys for different languages
        if section_name in ['title', 'título', 'titre']:
            section_name = 'title'
        elif section_name in ['abstract', 'resumen', 'résumé']:
            section_name = 'abstract'
        elif section_name in ['index', 'índice', 'indice']:
            section_name = 'index'
        
        section_content = match.group(2).strip()
        sections[section_name] = section_content
    return sections


def extract_points_and_subpoints(index_text):
    normalized_text = index_text.replace('\r\n', '\n').replace('\r', '\n').strip()

    # Pattern to match main sections (Roman numerals or numerical)
    main_section_pattern = re.compile(r'^((?:[IVXLCDM]+|\d+))\.\s+(.+)', re.MULTILINE)
    # Pattern to match subsections (Alphabetic or numerical following a main section identifier)
    subsection_pattern = re.compile(r'^\s{3,}((?:[A-Z]|\d+\.\d+))\.\s+(.+)', re.MULTILINE)
    # Pattern to match subsubsections (Numerical following a subsection identifier, e.g., "1.1.1.")
    subsubsection_pattern = re.compile(r'^\s{6,}(\d+\.\d+\.\d+)\.\s+(.+)', re.MULTILINE)

    results = []

    for main_section in main_section_pattern.finditer(normalized_text):
        main_text = main_section.group(2).strip()
        results.append((main_section.group(1), main_text))

        section_end = None
        next_main = main_section_pattern.search(normalized_text, main_section.end())
        if next_main:
            section_end = next_main.start()

        section_text = normalized_text[main_section.start():section_end]

        for subsection in subsection_pattern.finditer(section_text):
            sub_text = subsection.group(2).strip()
            results.append((f"{main_section.group(1)}.{subsection.group(1)}", sub_text))

            subsection_end = None
            next_sub = subsection_pattern.search(section_text, subsection.end())
            if next_sub:
                subsection_end = next_sub.start()

            sub_section_text = section_text[subsection.start():subsection_end]

            for subsubsection in subsubsection_pattern.finditer(sub_section_text):
                subsub_text = subsubsection.group(2).strip()
                results.append((f"{main_section.group(1)}.{subsection.group(1)}.{subsubsection.group(1)}", subsub_text))

    return results


def generate_paragraphs(index, abstract, documents, inverted_index, length, language):
    paragraphs = []
    references = {}
    
    points_and_subpoints = extract_points_and_subpoints(index)

    headings = []
    
    for section_number, point in points_and_subpoints:
        full_section = f"{section_number}: {point}"
        if '.' in section_number:
            full_section = f"{headings[-1]} \t {section_number}: {point}"
            related_terms = llm_router.generate("claude-3-haiku-20240307", [{"role": "user", "content": f"Generate 3 related terms for the following topic: {point}"}],
                                                    max_tokens=50, temperature=0.6, top_p=1.0, system="You are an AI assistant that helps generate related terms for a given topic. The format must be 'term1, term2, term3'. Separate each terms with commas and do no write anything else but the terms.")
                
            related_terms = related_terms.split(", ")
            
            relevant_documents = []

            for term in related_terms:
                print("Buscando...")
                relevant_documents = search(term, inverted_index, documents)

            references[section_number] = relevant_documents
            
            num_paragraphs = {"very short": 1, "short": 3, "medium": 5, "long": 8}[length]
            
            for i in range(num_paragraphs):
                if i == 0:
                    messages = [
                        {"role": "user", "content": f"Write a paragraph for a paper based on the following information:\n\nIndex: {index}\nAbstract: {abstract}\nRelevant documents: {relevant_documents}\nSpecific section: {full_section}\n\nRemember to focus on the specific section and not expand upon other parts of the index. There are {num_paragraphs - i} paragraphs left to write for this section."}
                    ]
                else:
                    messages = [
                        {"role": "user", "content": f"Write a paragraph for a paper based on the following information:\n\nIndex: {index}\nAbstract: {abstract}\nRelevant documents: {relevant_documents}\nAlready written: {paragraphs}\nSpecific section: {full_section}\n\nRemember to focus on the specific section and not expand upon other parts of the index. There are {num_paragraphs - i} paragraphs left to write for this section."}
                    ]

                paragraph = llm_router.generate("gpt-4-turbo-preview", messages, max_tokens=400, temperature=0.85, top_p=0.95, system=f"You are an expert professor AI with a formed mind and opinions that writes paragraphs for an academic paper, you are able to understand and write about complex topics in an academic manner with technical language in perfect {language}, as would be seen from a doctorate. Write the paragraph in markdown format. Your paragraph must be written as fully integrated in the text, do not mention anything about its structure nor metarreference anything outside of it.")
                print(paragraph)
                paragraphs.append((section_number, paragraph))
        else:
            headings.append(full_section)

    return paragraphs, references

def generate_citations(paper, references, citation_style, language):
    messages = [
        {"role": "user", "content": f"Generate citations for the following paper in {citation_style} style. Even if they are repeated, only mention each work once. The references for each section are provided as a dictionary:\n\nPaper:\n{paper}\n\nReferences:\n{json.dumps(references)}"}
    ]
    
    citations = llm_router.generate("claude-3-sonnet-20240229", messages, max_tokens=850, temperature=0.5, top_p=1.0, system=f"You are an AI assistant specialized in writing citations.")
    
    return citations

def generate_title(index, language):
    messages = [
        {"role": "user", "content": f"Generate a title for an essay based on the following index:\n\n{index}"}
    ]
    
    title = llm_router.generate("claude-3-opus-20240229", messages, max_tokens=100, temperature=0.95, top_p=0.9, system=f"You are an expert AI professor with a formed mind and opinions that specializes in generating titles for academic papers, you are able to understand and write about complex topics in an academic manner with technical language in perfect {language}, as would be seen from a doctorate. You must only generate the title and nothing more, limit yourself to that, do not include anything other than the tile. That is, you must not include 'Title:' or any other variant, simply the plain title.")
    
    return title

def convert_to_pdf(paper):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=3*cm, rightMargin=3*cm, topMargin=3*cm, bottomMargin=3*cm)
    styles = getSampleStyleSheet()
    
    normal_style = ParagraphStyle(
        'Normal',
        fontName='Times-Roman',
        fontSize=12,
        leading=24,  # Double line spacing
        firstLineIndent=1.25*cm,  # First line indent of 1.25 cm
        alignment=4  # Justified text
    )

    # Update styles for headings
    styles['Heading1'].fontName = 'Times-Bold'
    styles['Heading1'].fontSize = 18
    styles['Heading1'].alignment = 1  # Center alignment
    styles['Heading2'].fontName = 'Times-Bold'
    styles['Heading2'].fontSize = 16
    styles['Heading2'].alignment = 1  # Center alignment

    elements = []

    # Use BeautifulSoup to ensure HTML tags from Markdown are properly closed
    md = markdown.Markdown(extensions=['extra', 'nl2br', 'sane_lists'])
    html = md.convert(paper)
    soup = BeautifulSoup(html, 'html.parser')

    for tag in soup.find_all(True):
        text = str(tag.string)
        if tag.name == 'h1':
            elements.append(Paragraph(text, styles['Heading1']))
        elif tag.name == 'h2':
            elements.append(Paragraph(text, styles['Heading2']))
        elif tag.name in ['p', 'li']:  # Handle paragraphs and list items similarly
            elements.append(Paragraph(text, normal_style))
        elements.append(Spacer(1, 12))

    doc.build(elements)
    return buffer.getvalue()

def convert_to_docx(paper):
    doc = Document()
    
    md = markdown.Markdown()
    html = md.convert(paper)
    
    for element in md.parser.elements:
        if element.tag == 'p':
            doc.add_paragraph(element.text)
        elif element.tag.startswith('h'):
            level = int(element.tag[1])
            doc.add_heading(element.text, level=level)
        elif element.tag == 'image':
            doc.add_picture(element.attrib['src'])
    
    buffer = io.BytesIO()
    doc.save(buffer)
    
    return buffer.getvalue()

def main():
    st.title("Academic Essay Generator")
    
    instruction = st.text_input("Enter your essay instructions")
    length = st.selectbox("Select essay length", ["very short", "short", "medium", "long"])
    language = st.selectbox("Select language", ["English", "Spanish", "French"])
    citation_style = st.selectbox("Select citation style", ["APA", "Chicago"])
    
    if st.button("Generate Essay"):
        with st.spinner("Generating index and abstract..."):
            index_and_abstract = generate_index_and_abstract(instruction, length, language)
            print(index_and_abstract)
        
        queries = llm_router.generate("claude-3-haiku-20240307", [{"role": "user", "content": f"Generate 5 different search queries based on the following index and abstract in {language}:\n\n{index_and_abstract}"}],
                                      max_tokens=200, temperature=0.6, top_p=1.0, system="You are an AI assistant that helps generate related terms to search academic papers based on this index and abstract. The format must be 'term1, term2, term3'. Separate each terms with commas and do no write anything else but the terms.")
        print(queries)
        
        with st.spinner("Downloading relevant papers..."):
            print("Downloading relevant papers...")
            download_papers(queries)
            print("Downloaded")
        
        with st.spinner("Processing downloaded papers..."):
            print("Processing downloaded papers...")
            documents, inverted_index = process_downloaded_papers()
        
        with st.spinner("Generating essay paragraphs..."):
            sections = extract_sections(index_and_abstract)
            print(sections)
            paragraphs, references = generate_paragraphs(sections.get('index', ''), sections.get('abstract', ''), documents, inverted_index, length, language)

        with st.spinner("Generating citations..."):
            citations = generate_citations(paragraphs, references, citation_style, language)

        title = generate_title((sections.get('index', "") + sections.get('abstract', '')), language)
        
        paper = f"# {title}\n\n## Abstract\n{sections.get('abstract', '')}\n\n## Index\n"

        # Add index with preserved formatting
        index_lines = sections.get('index', '').split('\n')
        for line in index_lines:
            indent_level = len(line) - len(line.lstrip())
            if indent_level >= 6:  # Assuming subsubsections are indented by at least 6 spaces
                paper += f"    - {line.strip()}\n"
            elif indent_level > 0:  # Assuming subsections are indented but less than subsubsections
                paper += f"  - {line.strip()}\n"
            else:  # Main sections with no indentation
                paper += f"- {line.strip()}\n"

        last_main_section = None  # Keep track of the last main section number added to the paper

        for section_number, paragraph in paragraphs:
            main_section_part = section_number.split('.')[0]  # Extract the main section part
            
            # Check if this is the main section and it's different from the last one added
            if section_number.count('.') == 0 and main_section_part != last_main_section:
                last_main_section = main_section_part  # Update the last main section
                paper += f"\n## {section_number} \n{paragraph}\n"
            # For subsections and subsubsections, don't add the main section number again
            elif section_number.count('.') == 1:
                paper += f"\n### {section_number} \n{paragraph}\n"
            else:
                paper += f"\n#### {section_number} \n{paragraph}\n"

        # Format citations with each citation on a new line
        citation_lines = citations.strip().split('\n')  # Assuming each citation is separated by a newline
        formatted_citations = '\n'.join([f"- {line.strip()}" for line in citation_lines if line.strip()])  # Prepend '- ' to each citation for Markdown list formatting

        paper += f"\n## References\n{formatted_citations}"

        st.markdown(paper)
        
        pdf_data = convert_to_pdf(paper)
        # docx_data = convert_to_docx(paper)
        
        st.download_button("Download PDF", data=pdf_data, file_name=f"{title}.pdf", mime="application/pdf")
        # st.download_button("Download DOCX", data=docx_data, file_name="essay.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    main()